import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement, parse_xml # qn needed for OxmlElement attribute setting if used directly, not here.
import os

# --- Styling Utilities ---
def set_cell_background(cell, color_hex):
    """Set background color of a cell (color_hex: e.g. 'c3b249')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_run_style(run, font_name, font_size, color=None, bold=False, underline=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.underline = underline
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_hyperlink(paragraph, url, text, color_hex='0000FF', font_name='Roboto', font_size=12):
    # python-docx does not support hyperlinks natively, so we use XML hack
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), color_hex)
    rPr.append(color)
    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    # Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    # Size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

def load_json_data(json_path):
    """Loads data from a JSON file."""
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)

# From your code.py
def replace_text_in_paragraph(paragraph, old_text, new_text):
    """Replaces text in a paragraph, handling cases where the text might be split across runs."""
    if old_text in paragraph.text:
        # Store initial runs since we'll be modifying them
        initial_runs = list(paragraph.runs)
        
        # Find the runs containing parts of old_text
        start_run_index = None
        end_run_index = None
        accumulated_text = ""
        
        for i, run in enumerate(initial_runs):
            accumulated_text += run.text
            if old_text in accumulated_text:
                end_run_index = i
                # Find start by working backwards
                temp_text = ""
                for j in range(i, -1, -1):
                    temp_text = initial_runs[j].text + temp_text
                    if old_text in temp_text:
                        start_run_index = j
                        break
                break
        
        if start_run_index is not None and end_run_index is not None:
            # Get the text before and after the replacement in the affected runs
            full_text = "".join(run.text for run in initial_runs[start_run_index:end_run_index + 1])
            start_pos = full_text.find(old_text)
            end_pos = start_pos + len(old_text)
            
            before_text = full_text[:start_pos]
            after_text = full_text[end_pos:]
            
            # Replace the content in the first run
            initial_runs[start_run_index].text = before_text + str(new_text) + after_text
            
            # Remove the other runs that were part of the placeholder
            for i in range(start_run_index + 1, end_run_index + 1):
                initial_runs[i].text = ""
            
            # Remove bold for date values
            if old_text in ["{{DATE_OF_SITE_AUDIT}}", "{{DATE_OF_REPORT_TO_CLIENT}}"]:
                initial_runs[start_run_index].bold = False

def replace_text_in_shapes(doc, replacements):
    """Replaces text placeholders inside shapes in the document."""
    for shape in doc.element.xpath('//w:drawing//w:t'):
        text = shape.text
        if text:
            for key, value in replacements.items():
                if key in text:
                    shape.text = text.replace(key, str(value))
                    # Remove bold for date values
                    if key in ["{{DATE_OF_SITE_AUDIT}}", "{{DATE_OF_REPORT_TO_CLIENT}}"]:
                        shape.bold = False

def replace_global_placeholders(doc, replacements):
    """Replaces global text placeholders throughout the document."""
    for key, value in replacements.items():
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, key, value)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, key, value)
    replace_text_in_shapes(doc, replacements)

def find_placeholder_paragraph(doc, placeholder_text):
    """Finds the first paragraph object containing the exact placeholder_text."""
    for p in doc.paragraphs:
        if placeholder_text in p.text:
            return p
    return None

def remove_paragraph_element(paragraph_element):
    """Removes a paragraph element from its parent."""
    if paragraph_element is not None and paragraph_element.getparent() is not None:
        paragraph_element.getparent().remove(paragraph_element)

def add_table_borders(table):
    """Adds borders to a table."""
    tbl = table._tbl
    tblPr = tbl.get_or_add_tblPr()
    tblBorders = OxmlElement('w:tblBorders')
    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border_element = OxmlElement(f'w:{border}')
        border_element.set(qn('w:val'), 'single') # Use qn for attribute names with namespace prefix
        border_element.set(qn('w:sz'), '4')
        border_element.set(qn('w:space'), '0')
        border_element.set(qn('w:color'), '000000')
        tblBorders.append(border_element)
    tblPr.append(tblBorders)

def add_border_to_paragraph(paragraph):
    """Add a border around a paragraph"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    # Add borders on all sides
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')  # single line border
        border.set(qn('w:sz'), '24')  # border width
        border.set(qn('w:space'), '1')  # space between border and text
        border.set(qn('w:color'), '000000')  # black color
        pBdr.append(border)
    
    pPr.append(pBdr)

def add_border_to_run(run):
    """Add border to a specific run within a paragraph"""
    rPr = run._r.get_or_add_rPr()
    bdr = OxmlElement('w:bdr')
    bdr.set(qn('w:val'), 'single')  # single line border
    bdr.set(qn('w:sz'), '8')  # border width
    bdr.set(qn('w:space'), '6')  # space between border and text
    bdr.set(qn('w:color'), '000000')  # black color
    rPr.append(bdr)

def create_site_number_para(placeholder_para, site_number):
    """Creates a paragraph with bordered site number"""
    site_num_para = placeholder_para.insert_paragraph_before()
    site_text = f"SITE # {site_number}"
    run = site_num_para.add_run(site_text)
    set_run_style(run, 'Roboto', 12)
    add_border_to_run(run)
    site_num_para.paragraph_format.space_before = Pt(6)
    site_num_para.paragraph_format.space_after = Pt(6)
    return site_num_para

# --- Section Handler Functions ---

def handle_detailed_report_table(doc, placeholder_para, section_config):
    """Handles the detailed report table section."""
    if not (section_config and section_config.get("render") and section_config.get("data")):
        if placeholder_para: remove_paragraph_element(placeholder_para._element)
        return

    data = section_config["data"]
    num_cols = 7
    table = doc.add_table(rows=1, cols=num_cols) # Table object created but not yet in main doc body
    table.style = doc.styles['Table Grid']

    # Agency scope row
    agency_scope_text = f"Agency Name: {data['agency_name']} Scope: ({data['scope_total_urls']}) URLs Compliant Status: {data['overall_compliant_status']}"
    hdr_cells_agency = table.rows[0].cells
    merged_cell_agency = hdr_cells_agency[0].merge(hdr_cells_agency[num_cols-1])
    merged_cell_agency.text = ''  # Clear first
    para = merged_cell_agency.paragraphs[0]
    para.clear()
    # Compose styled runs for each part
    agency_name = str(data['agency_name'])
    scope_urls = f"({data['scope_total_urls']}) URLs"
    compliant_status = str(data['overall_compliant_status'])
    before = f"Agency Name: "
    after = f"\t\t\tScope: "
    between = f"\t\t\tCompliant Status: "
    # Add styled runs with bold
    run = para.add_run(before)
    set_run_style(run, 'Roboto', 12, bold=True)
    run2 = para.add_run(agency_name)
    set_run_style(run2, 'Roboto', 12, color='E36C09', bold=True)
    run3 = para.add_run(after)
    set_run_style(run3, 'Roboto', 12, bold=True)
    run4 = para.add_run(scope_urls)
    set_run_style(run4, 'Roboto', 12, color='0000FF', bold=True)
    run5 = para.add_run(between)
    set_run_style(run5, 'Roboto', 12, bold=True)
    run6 = para.add_run(compliant_status)
    if compliant_status.strip().lower() == 'compliant':
        set_run_style(run6, 'Roboto', 12, color='006050', bold=True)
    else:
        set_run_style(run6, 'Roboto', 12, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_cell_background(merged_cell_agency, 'f2dcdb')

    # Header row
    headers = ["Site", "URL details", "Date of site visit", "Has pay portal?\n(Y/N)",
               "If yes, is site PCI-compliant?\n(Y/N)", "If not PCI-compliant,\ndate reported to DFA",
               "Notes / Observations"]
    hdr_row_cols = table.add_row().cells
    for i, header_text in enumerate(headers):
        cell_paragraph = hdr_row_cols[i].paragraphs[0]
        cell_paragraph.clear()
        run = cell_paragraph.add_run(header_text)
        set_run_style(run, 'Roboto', 12, bold=True)
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_row_cols[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(hdr_row_cols[i], 'c3b249')

    # Data rows
    for item in data["urls"]:
        row_cells = table.add_row().cells
        # Site serial
        para0 = row_cells[0].paragraphs[0]
        para0.clear()
        run0 = para0.add_run(str(item.get("site_serial", "")))
        set_run_style(run0, 'Roboto', 12)
        # URL details (hyperlink if looks like URL)
        para1 = row_cells[1].paragraphs[0]
        para1.clear()
        url_val = item.get("url_details", "")
        if url_val.startswith('http://') or url_val.startswith('https://'):
            add_hyperlink(para1, url_val, url_val)
        else:
            run1 = para1.add_run(url_val)
            set_run_style(run1, 'Roboto', 12)
        # Date of site visit
        para2 = row_cells[2].paragraphs[0]
        para2.clear()
        run2 = para2.add_run(item.get("date_of_site_visit", ""))
        set_run_style(run2, 'Roboto', 12)
        # Has pay portal
        para3 = row_cells[3].paragraphs[0]
        para3.clear()
        run3 = para3.add_run(item.get("has_pay_portal", ""))
        set_run_style(run3, 'Roboto', 12)
        # PCI compliant
        para4 = row_cells[4].paragraphs[0]
        para4.clear()
        run4 = para4.add_run(item.get("is_pci_compliant", ""))
        set_run_style(run4, 'Roboto', 12)
        # Date reported to DFA
        para5 = row_cells[5].paragraphs[0]
        para5.clear()
        run5 = para5.add_run(item.get("date_reported_to_dfa", ""))
        set_run_style(run5, 'Roboto', 12)
        # Notes/Observations
        para6 = row_cells[6].paragraphs[0]
        para6.clear()
        run6 = para6.add_run(item.get("notes_observations", ""))
        set_run_style(run6, 'Roboto', 12)
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    if placeholder_para and placeholder_para._p.getparent() is not None:
        # Insert the table's XML element before the placeholder paragraph's XML element
        placeholder_para._p.addprevious(table._tbl)
    
    if placeholder_para:
        # Remove original placeholder paragraph
        remove_paragraph_element(placeholder_para._element)

def _insert_pci_certified_content(placeholder_para, section_data, doc):
    # This helper function inserts content *before* placeholder_para
    data = section_data
    num_sites = len(data.get("sites", []))
    title_text = data.get("title_text_template", "URL(s) with Payment Option ({count}) – PCI Certified – NO ACTION REQUIRED:").format(count=num_sites)
    heading_para = placeholder_para.insert_paragraph_before()
    run = heading_para.add_run(title_text)
    set_run_style(run, 'Roboto', 12, color='0000FF', bold=True, underline=True)

    for site in data.get("sites", []):
        # Site number with border
        create_site_number_para(placeholder_para, site.get('site_number_display', ''))

        # URL line next
        url_para = placeholder_para.insert_paragraph_before()
        url_val = site['url']
        url_label = f"URL: "
        run_label = url_para.add_run(url_label)
        set_run_style(run_label, 'Roboto', 12)
        if url_val.startswith('http://') or url_val.startswith('https://'):
            add_hyperlink(url_para, url_val, url_val)
        else:
            run_url = url_para.add_run(url_val)
            set_run_style(run_url, 'Roboto', 12, color='0000FF')

        for screenshot in site.get("screenshots", []):
            desc_para = placeholder_para.insert_paragraph_before(screenshot['description'])
            for run in desc_para.runs:
                set_run_style(run, 'Roboto', 12)
            img_p = placeholder_para.insert_paragraph_before()
            if os.path.exists(screenshot['image_path']):
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.add_run().add_picture(screenshot['image_path'], width=Inches(6.0))
            else:
                img_p.add_run(f"[Image not found: {screenshot['image_path']}]")
        
        # Create a single paragraph for the entire conclusion section
        concl_para = placeholder_para.insert_paragraph_before()
        
        # Add conclusion intro
        concl_intro = site.get('conclusion_intro', "Conclusion:")
        run_intro = concl_para.add_run(concl_intro + "\n")
        set_run_style(run_intro, 'Roboto', 12)
        
        # Add conclusion text
        concl_text = site['conclusion_text'] + "\n\n"
        run_text = concl_para.add_run(concl_text)
        set_run_style(run_text, 'Roboto', 12)
        
        # Add action status in the same paragraph
        action_status = site['action_status']
        run_status = concl_para.add_run(action_status)
        set_run_style(run_status, 'Roboto', 12)
        
        # Add border to the entire conclusion paragraph
        add_border_to_paragraph(concl_para)
        
        # Remove original placeholder paragraph
        remove_paragraph_element(placeholder_para._element)

def handle_pci_certified_sites_section(doc, placeholder_para, section_config):
    if not (section_config and section_config.get("render") and section_config.get("data")):
        if placeholder_para: remove_paragraph_element(placeholder_para._element)
        return
    
    _insert_pci_certified_content(placeholder_para, section_config["data"], doc)

    if placeholder_para:
        # Remove original placeholder paragraph
        remove_paragraph_element(placeholder_para._element)

def _insert_clarification_needed_content(placeholder_para, section_data, doc):
    # This helper function inserts content *before* placeholder_para
    data = section_data
    num_sites = len(data.get("sites", []))
    title_text = data.get("title_text_template", "URL(s) with Payment Option ({count}) – To be clarified by the Agency:").format(count=num_sites)
    heading_para = placeholder_para.insert_paragraph_before()
    run = heading_para.add_run(title_text)
    set_run_style(run, 'Roboto', 12, color='0000FF', bold=True, underline=True)

    for site in data.get("sites", []):
        # Site number with border
        create_site_number_para(placeholder_para, site.get('site_number_display', ''))

        # URL line next
        url_para = placeholder_para.insert_paragraph_before()
        url_val = site.get('site_main_url', '')
        url_label = f"URL: "
        run_label = url_para.add_run(url_label)
        set_run_style(run_label, 'Roboto', 12)
        if url_val.startswith('http://') or url_val.startswith('https://'):
            add_hyperlink(url_para, url_val, url_val)
        else:
            run_url = url_para.add_run(url_val)
            set_run_style(run_url, 'Roboto', 12, color='0000FF')

        for scenario in site.get("scenarios", []):
            scenario_para = placeholder_para.insert_paragraph_before(scenario.get('scenario_label', 'Scenario:'))
            for run in scenario_para.runs:
                set_run_style(run, 'Roboto', 12, bold=True)
            for step in scenario.get("steps", []):
                desc_para = placeholder_para.insert_paragraph_before(step['description'])
                for run in desc_para.runs:
                    set_run_style(run, 'Roboto', 12)
                img_p = placeholder_para.insert_paragraph_before()
                if os.path.exists(step['image_path']):
                    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_p.add_run().add_picture(step['image_path'], width=Inches(6.0))
                else:
                    img_p.add_run(f"[Image not found: {step['image_path']}]")
            # Add ACTION REQUIRED text without border
            action_intro_para = placeholder_para.insert_paragraph_before(scenario.get("action_required_intro", "ACTION REQUIRED:"))
            for run in action_intro_para.runs:
                set_run_style(run, 'Roboto', 12, bold=True)
            
            # Get the template and format it with the data
            template = scenario.get("action_text_template", (
                'The URL {site_main_url} links to "{third_party_url_in_question}" for payments. '
                'This site was not in scope for {agency_abbreviation}. '
                'Please clarify with {contact_person_name} ({contact_person_email}): '
                'SITE NUMBER: {site_number_display} '
                'SITE URL: {site_main_url} '
                'a. Is "{third_party_url_in_question}" owned/managed by {agency_abbreviation}? '
                'If "Yes", provide relationship details. '
                'If "No", provide details (e.g., convenience link). This site will be considered out of scope.'
            ))
            
            # Format the template with the data
            content = template.format(
                site_main_url=site.get("site_main_url", ""),
                third_party_url_in_question=scenario.get("third_party_url_in_question", ""),
                agency_abbreviation=site.get("agency_abbreviation", ""),
                contact_person_name=data.get("contact_person_name", ""),
                contact_person_email=data.get("contact_person_email", ""),
                site_number_display=site.get("site_number_display", "")
            )
            
            # Create paragraph and preserve newlines
            content_para = placeholder_para.insert_paragraph_before()
            # Set left alignment explicitly
            content_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Split content by newlines and add each part
            parts = content.split('\n')
            for i, part in enumerate(parts):
                if i > 0:  # Add newline before all parts except the first
                    content_para.add_run('\n')
                run = content_para.add_run(part)
                set_run_style(run, 'Roboto', 12)

def handle_clarification_needed_sites_section(doc, placeholder_para, section_config):
    """Handles the 'clarification needed sites' section."""
    if not (section_config and section_config.get("render") and section_config.get("data")):
        if placeholder_para: remove_paragraph_element(placeholder_para._element)
        return
        
    _insert_clarification_needed_content(placeholder_para, section_config["data"], doc)

    if placeholder_para:
        # Remove original placeholder paragraph
        remove_paragraph_element(placeholder_para._element)

def handle_access_error_sites_section(doc, placeholder_para, section_config):
    """Handles the 'access error sites' section."""
    if not (section_config and section_config.get("render") and section_config.get("data")):
        if placeholder_para: remove_paragraph_element(placeholder_para._element)
        return

    data = section_config["data"]
    num_sites = len(data.get("sites", []))
    title_text = data.get("title_text_template", "URLs with Access Errors ({count}) – ACTION REQUIRED:").format(count=num_sites)
    heading_para = placeholder_para.insert_paragraph_before()
    run = heading_para.add_run(title_text)
    set_run_style(run, 'Roboto', 12, color='0000FF', bold=True, underline=True)

    for site in data.get("sites", []):
        # Site number with border
        create_site_number_para(placeholder_para, site.get('site_number_display', ''))

        # URL line next
        url_para = placeholder_para.insert_paragraph_before()
        url_val = site.get('url', '')
        url_label = f"URL: "
        run_label = url_para.add_run(url_label)
        set_run_style(run_label, 'Roboto', 12)
        if url_val.startswith('http://') or url_val.startswith('https://'):
            add_hyperlink(url_para, url_val, url_val)
        else:
            run_url = url_para.add_run(url_val)
            set_run_style(run_url, 'Roboto', 12, color='0000FF')

        # Add screenshot first
        img_p = placeholder_para.insert_paragraph_before()
        if os.path.exists(site.get('screenshot_path', '')):
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_p.add_run().add_picture(site.get('screenshot_path'), width=Inches(6.0))
        else:
            img_p.add_run(f"[Image not found: {site.get('screenshot_path', '')}]")
        
        # Then add ACTION REQUIRED text and content
        action_intro_para = placeholder_para.insert_paragraph_before(site.get('action_required_intro', 'ACTION REQUIRED:'))
        for run in action_intro_para.runs:
            set_run_style(run, 'Roboto', 12, bold=True)
        
        error_desc_para = placeholder_para.insert_paragraph_before(site.get('error_description', ''))
        for run in error_desc_para.runs:
            set_run_style(run, 'Roboto', 12)
        
        action_points_prompt_para = placeholder_para.insert_paragraph_before(site.get('action_points_prompt', ''))
        for run in action_points_prompt_para.runs:
            set_run_style(run, 'Roboto', 12)
        
        for point in site.get("action_points", []):
            point_para = placeholder_para.insert_paragraph_before(f"☐ {point}")
            for run in point_para.runs:
                set_run_style(run, 'Roboto', 12)

    if placeholder_para:
        # Remove original placeholder paragraph
        remove_paragraph_element(placeholder_para._element)

def handle_login_error_sites_section(doc, placeholder_para, section_config):
    """Handles the 'login error sites' section."""
    if not (section_config and section_config.get("render") and section_config.get("data")):
        if placeholder_para: remove_paragraph_element(placeholder_para._element)
        return

    data = section_config["data"]
    num_sites = len(data.get("sites", []))
    title_text = data.get("title_text_template", "URLs with Login Errors ({count}) – ACTION REQUIRED:").format(count=num_sites)
    heading_para = placeholder_para.insert_paragraph_before()
    run = heading_para.add_run(title_text)
    set_run_style(run, 'Roboto', 12, color='0000FF', bold=True, underline=True)

    for site in data.get("sites", []):
        # Site number with border
        create_site_number_para(placeholder_para, site.get('site_number_display', ''))

        # URL line next
        url_para = placeholder_para.insert_paragraph_before()
        url_val = site.get('url', '')
        url_label = f"URL: "
        run_label = url_para.add_run(url_label)
        set_run_style(run_label, 'Roboto', 12)
        if url_val.startswith('http://') or url_val.startswith('https://'):
            add_hyperlink(url_para, url_val, url_val)
        else:
            run_url = url_para.add_run(url_val)
            set_run_style(run_url, 'Roboto', 12, color='0000FF')

        # Add screenshot first
        img_p = placeholder_para.insert_paragraph_before()
        if os.path.exists(site.get('screenshot_path', '')):
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_p.add_run().add_picture(site.get('screenshot_path'), width=Inches(6.0))
        else:
            img_p.add_run(f"[Image not found: {site.get('screenshot_path', '')}]")
        
        # Then add ACTION REQUIRED text and content
        action_intro_para = placeholder_para.insert_paragraph_before(site.get('action_required_intro', 'ACTION REQUIRED:'))
        for run in action_intro_para.runs:
            set_run_style(run, 'Roboto', 12, bold=True)
        
        error_desc_para = placeholder_para.insert_paragraph_before(site.get('error_description', ''))
        for run in error_desc_para.runs:
            set_run_style(run, 'Roboto', 12)
        
        action_points_prompt_para = placeholder_para.insert_paragraph_before(site.get('action_points_prompt', ''))
        for run in action_points_prompt_para.runs:
            set_run_style(run, 'Roboto', 12)
        
        for point in site.get("action_points", []):
            point_para = placeholder_para.insert_paragraph_before(f"☐ {point}")
            for run in point_para.runs:
                set_run_style(run, 'Roboto', 12)

    if placeholder_para:
        # Remove original placeholder paragraph
        remove_paragraph_element(placeholder_para._element)

# --- Main Generator ---
def generate_report(template_path, json_path, output_path):
    """Generates the DOCX report."""
    if not os.path.exists(template_path):
        print(f"Error: Template file '{template_path}' not found. Please provide a valid template.")
        return
    if not os.path.exists(json_path):
        print(f"Error: JSON data file '{json_path}' not found.")
        return

    try:
        doc = Document(template_path)
        data = load_json_data(json_path)
    except Exception as e:
        print(f"Error loading template or JSON data: {e}")
        return

    if "global_text_replacements" in data:
        replacements = data["global_text_replacements"]
        if "{{OBSERVATIONS}}" in replacements and isinstance(replacements.get("{{OBSERVATIONS}}"), list):
            formatted_observations = "\n".join(f"• {item}" for item in replacements["{{OBSERVATIONS}}"])
            replacements["{{OBSERVATIONS}}"] = formatted_observations
        replace_global_placeholders(doc, replacements)
    


    
    sections_map = {
        "detailed_report_table": handle_detailed_report_table,
        "pci_certified_sites": handle_pci_certified_sites_section,
        "clarification_needed_sites": handle_clarification_needed_sites_section,
        "access_error_sites": handle_access_error_sites_section,
        "Login_error_sites": handle_login_error_sites_section
    }

    if "sections" in data:
        first_section = True
        for section_key, section_config in data["sections"].items():
            placeholder_text = section_config.get("placeholder")
            if not placeholder_text:
                print(f"Warning: No placeholder text defined for section '{section_key}'. Skipping.")
                continue
            
            placeholder_para = find_placeholder_paragraph(doc, placeholder_text)
            
            if placeholder_para:
                handler_func = sections_map.get(section_key)
                if handler_func:
                    print(f"Processing section: {section_key} with placeholder: '{placeholder_text}'")
                    try:
                        # Insert page break before section, except for the first section
                        if not first_section:
                            pb_para = placeholder_para.insert_paragraph_before()
                            pb_para.add_run().add_break(WD_BREAK.PAGE)
                        handler_func(doc, placeholder_para, section_config)
                        first_section = False
                    except Exception as e_handler:
                        print(f"Error processing section '{section_key}': {e_handler}")
                        # Optionally, decide if placeholder should be removed or an error message inserted
                else:
                    print(f"Warning: No handler function defined for section '{section_key}'.")
                    # If no handler, decide if placeholder should be removed if not rendering
                    if not section_config.get("render", True) and placeholder_para:
                         remove_paragraph_element(placeholder_para._element)
            elif section_config.get("render", False):
                 print(f"Warning: Placeholder '{placeholder_text}' for section '{section_key}' not found in template, but section is set to render.")

    try:
        doc.save(output_path)
        print(f"Report generated successfully: {output_path}")
    except Exception as e:
        print(f"Error saving document: {e}")

# Helper to get qn for OxmlElement attributes if not already available
from docx.oxml.ns import qn

if __name__ == "__main__":
    template_file = "C:\\Users\\Harshini_S\\Documents\\GitHub\\Matrix\\template.docx" # Ensure this path is correct
    json_data_file = "C:\\Users\\Harshini_S\\Documents\\GitHub\\Matrix\\data.json"  # Ensure this path is correct
    output_file = "C:\\Users\\Harshini_S\\Documents\\GitHub\\Matrix\\Generated_Report.docx"

    generate_report(template_file, json_data_file, output_file)